import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.api_clients import _batch_translate, _build_adaptive_batches, legacy_batch_translate
from core.constants import FINALIZATION_PROGRESS_STEPS, GEMINI_REQUEST_CHAR_LIMIT
from core.progress import TranslationProgress
from core.run_fusion import fuse_merged_paragraph_runs, merge_adjacent_runs_inplace
from core.text_utils import normalize_run_tags
from core.utils import _get_gemini_execution_settings, _is_gemini_backend, _summarize_numeric_values


# ─── RTL / Arabic Formatting ─────────────────────────────────

def _apply_rtl_formatting(run, para):
    pPr = para._element.get_or_add_pPr()
    rPr = run._element.get_or_add_rPr()

    bidi_tag = qn('w:bidi')
    if pPr.find(bidi_tag) is None:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')

        pPr_after_tags = {
            qn('w:adjustRightInd'), qn('w:snapToGrid'), qn('w:spacing'), qn('w:ind'),
            qn('w:contextualSpacing'), qn('w:mirrorIndents'), qn('w:suppressOverlap'),
            qn('w:jc'), qn('w:textDirection'), qn('w:textAlignment'),
            qn('w:textboxTightWrap'), qn('w:outlineLvl'), qn('w:divId'),
            qn('w:cnfStyle'), qn('w:rPr'), qn('w:sectPr'), qn('w:pPrChange')
        }

        inserted = False
        for child in pPr:
            if child.tag in pPr_after_tags:
                child.addprevious(bidi)
                inserted = True
                break
        if not inserted:
            pPr.append(bidi)

    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run.font.rtl = True
    rtl_tag = qn('w:rtl')
    if rPr.find(rtl_tag) is None:
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')

        rPr_after_tags_rtl = {
            qn('w:cs'), qn('w:em'), qn('w:lang'), qn('w:eastAsianLayout'),
            qn('w:specVanish'), qn('w:oMath')
        }

        inserted = False
        for child in rPr:
            if child.tag in rPr_after_tags_rtl:
                child.addprevious(rtl)
                inserted = True
                break
        if not inserted:
            rPr.append(rtl)

    lang_tag = qn('w:lang')
    lang_element = rPr.find(lang_tag)

    if lang_element is None:
        lang_element = OxmlElement('w:lang')
        rPr_after_tags_lang = {
            qn('w:eastAsianLayout'), qn('w:specVanish'), qn('w:oMath')
        }

        inserted = False
        for child in rPr:
            if child.tag in rPr_after_tags_lang:
                child.addprevious(lang_element)
                inserted = True
                break
        if not inserted:
            rPr.append(lang_element)

    lang_element.set(qn('w:bidi'), 'ar-SA')


# ─── Paragraph Utilities ─────────────────────────────────────

def _get_run_style_dict(run):
    style = {}
    if run.font.name: style['font'] = run.font.name
    if run.font.size: style['size'] = run.font.size.pt
    if run.bold: style['bold'] = True
    if run.italic: style['italic'] = True
    if run.underline: style['underline'] = True
    if run.font.color and run.font.color.rgb: style['color'] = str(run.font.color.rgb)
    if run.font.highlight_color: style['highlight'] = str(run.font.highlight_color)
    if run.font.superscript: style['script'] = 'super'
    if run.font.subscript: style['script'] = 'sub'
    return style


def _can_merge_adjacent_docx_runs(prev_run, run):
    """Merge when formatting matches and the current run contributes text."""
    return _get_run_style_dict(prev_run) == _get_run_style_dict(run) and run.text


def _paragraph_plain_text(paragraph) -> str:
    if not paragraph:
        return ""
    return re.sub(r"\s+", " ", "".join(run.text for run in paragraph.runs)).strip()


def _shorten_context_text(text: str, limit: int = 140) -> str:
    text = re.sub(r"</?r\d+>", "", text or "")
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


def _is_heading_paragraph(paragraph) -> bool:
    try:
        style_name = (paragraph.style.name or "").lower()
    except Exception:
        style_name = ""
    text = _paragraph_plain_text(paragraph)
    if not text:
        return False
    if "heading" in style_name or style_name in {"title", "subtitle"}:
        return True
    if len(text) <= 90 and text == text.upper() and any(ch.isalpha() for ch in text):
        return True

    words = re.findall(r"\b\w+\b", text)
    meaningful_runs = [run for run in paragraph.runs if (run.text or "").strip()]
    all_bold = bool(meaningful_runs) and all(run.bold for run in meaningful_runs)
    short_label = len(words) <= 8 and len(text) <= 90
    sentence_like = any(mark in text for mark in ".!?")

    if short_label and text.rstrip().endswith(":"):
        return True
    if short_label and all_bold and not sentence_like:
        return True

    return False


def _nearest_meaningful_text(texts, start_index, step):
    index = start_index + step
    while 0 <= index < len(texts):
        candidate = texts[index]
        if candidate and any(ch.isalpha() for ch in candidate):
            return candidate
        index += step
    return ""


def _build_paragraph_contexts(paragraphs, base_contexts=None):
    base_contexts = base_contexts or {}
    plain_texts = [_paragraph_plain_text(para) for para in paragraphs]
    section_labels = []
    current_section = ""

    for para, text in zip(paragraphs, plain_texts):
        if _is_heading_paragraph(para):
            current_section = text
        section_labels.append(current_section)

    paragraph_contexts = {}
    for index, para in enumerate(paragraphs):
        text = plain_texts[index]
        parts = []
        seen = set()

        base_value = base_contexts.get(id(para), "").strip()
        if base_value:
            parts.append(base_value)
            seen.add(base_value)

        section_value = section_labels[index]
        if section_value and section_value != text:
            value = f"Section: {section_value}"
            if value not in seen:
                parts.append(value)
                seen.add(value)

        is_short = len(text) < 90 or len(text.split()) <= 10
        if is_short:
            prev_text = _nearest_meaningful_text(plain_texts, index, -1)
            next_text = _nearest_meaningful_text(plain_texts, index, 1)
            if prev_text:
                value = f"Prev: {_shorten_context_text(prev_text, 110)}"
                if value not in seen:
                    parts.append(value)
                    seen.add(value)
            if next_text:
                value = f"Next: {_shorten_context_text(next_text, 110)}"
                if value not in seen:
                    parts.append(value)
                    seen.add(value)

        paragraph_contexts[id(para)] = " | ".join(parts)

    return paragraph_contexts


def _prepare_paragraph_units(paragraphs, paragraph_contexts):
    prepared_units = []

    for para in paragraphs:
        if not para.runs:
            continue

        merged_runs = merge_adjacent_runs_inplace(list(para.runs), _can_merge_adjacent_docx_runs)
        valid_runs, fused_para_text = fuse_merged_paragraph_runs(merged_runs)

        if not valid_runs:
            continue

        combined_text = "".join(core for _, _, _, _, core in valid_runs)
        if not combined_text.strip() or all(not ch.isalpha() for ch in combined_text):
            continue

        prepared_units.append(
            {
                "paragraph": para,
                "valid_runs": valid_runs,
                "fused_text": fused_para_text,
                "context": paragraph_contexts.get(id(para), ""),
            }
        )

    return prepared_units


# ─── Translation Application ─────────────────────────────────

def _apply_translated_paragraphs(prepared_units, translated_fused_texts, target_lang):
    for unit, new_fused_text in zip(prepared_units, translated_fused_texts):
        para = unit["paragraph"]
        valid_runs = unit["valid_runs"]
        new_fused_text = normalize_run_tags(new_fused_text)

        parsed_runs = {}
        for match in re.finditer(r"<r(\d+)>(.*?)</r\1>", new_fused_text, re.DOTALL):
            parsed_runs[int(match.group(1))] = match.group(2)

        if len(parsed_runs) == len(valid_runs):
            for idx, run, leading, trailing, _ in valid_runs:
                new_core = parsed_runs.get(idx, "")
                new_text = f"{leading}{new_core}{trailing}"
                if run.text != new_text:
                    run.text = new_text
                if target_lang == "Arabic":
                    _apply_rtl_formatting(run, para)
            continue

        clean_text = re.sub(r"</?r\d+>", "", new_fused_text)
        _, first_run, first_leading, _, _ = valid_runs[0]
        _, _, _, last_trailing, _ = valid_runs[-1]
        first_run.text = f"{first_leading}{clean_text.strip()}{last_trailing}"
        if target_lang == "Arabic":
            _apply_rtl_formatting(first_run, para)
        for _, run, _, _, _ in valid_runs[1:]:
            run.text = ""


def _translate_paragraphs(
    paragraphs,
    target_lang,
    backend,
    api_key,
    phrases,
    progress,
    domain_context="",
    paragraph_contexts=None,
    telemetry=None,
    set_total=True,
    temperature=None,
):
    if not paragraphs:
        return

    prepared_units = _prepare_paragraph_units(paragraphs, paragraph_contexts or {})
    if not prepared_units:
        return

    fused_texts = [unit["fused_text"] for unit in prepared_units]
    fused_contexts = [unit["context"] for unit in prepared_units]
    translated_fused_texts = _batch_translate(
        fused_texts,
        target_lang,
        backend,
        api_key,
        phrases,
        progress,
        domain_context,
        telemetry=telemetry,
        contexts=fused_contexts,
        set_total=set_total,
        temperature=temperature,
    )
    _apply_translated_paragraphs(prepared_units, translated_fused_texts, target_lang)


def _legacy_translate_paragraphs(
    paragraphs,
    target_lang,
    backend,
    api_key,
    phrases,
    progress,
    domain_context="",
    telemetry=None,
):
    if not paragraphs:
        return

    to_translate_paras = []
    fused_texts = []

    for para in paragraphs:
        if not para.runs:
            continue

        merged_runs = merge_adjacent_runs_inplace(list(para.runs), _can_merge_adjacent_docx_runs)
        valid_runs, fused_para_text = fuse_merged_paragraph_runs(merged_runs)

        if valid_runs:
            to_translate_paras.append((para, valid_runs))
            fused_texts.append(fused_para_text)

    if not fused_texts:
        return

    translated_fused_texts = legacy_batch_translate(
        fused_texts,
        target_lang,
        backend,
        api_key,
        phrases,
        progress,
        domain_context,
        telemetry=telemetry,
    )

    for (para, valid_runs), new_fused_text in zip(to_translate_paras, translated_fused_texts):
        new_fused_text = normalize_run_tags(new_fused_text)
        parsed_runs = {}
        for match in re.finditer(r"<r(\d+)>(.*?)</r\1>", new_fused_text, re.DOTALL):
            parsed_runs[int(match.group(1))] = match.group(2)

        if len(parsed_runs) == len(valid_runs):
            for idx, run, leading, trailing, _ in valid_runs:
                new_core = parsed_runs.get(idx, "")
                new_text = f"{leading}{new_core}{trailing}"
                if run.text != new_text:
                    run.text = new_text
                if target_lang == "Arabic":
                    _apply_rtl_formatting(run, para)
            continue

        clean_text = re.sub(r"</?r\d+>", "", new_fused_text)
        _, first_run, first_leading, _, _ = valid_runs[0]
        _, _, _, last_trailing, _ = valid_runs[-1]
        first_run.text = f"{first_leading}{clean_text.strip()}{last_trailing}"
        if target_lang == "Arabic":
            _apply_rtl_formatting(first_run, para)
        for _, run, _, _, _ in valid_runs[1:]:
            run.text = ""


# ─── Document Structure Walking ──────────────────────────────

def _find_textbox_paragraphs(element, parent) -> list:
    paragraphs = []
    for txbx in element.iter(qn("w:txbxContent")):
        for p_elem in txbx.iter(qn("w:p")):
            paragraphs.append(Paragraph(p_elem, parent))
    try:
        for vtxbx in element.iter(qn("v:textbox")):
            for p_elem in vtxbx.iter(qn("w:p")):
                paragraphs.append(Paragraph(p_elem, parent))
    except KeyError:
        pass

    unique = []
    seen = set()
    for p in paragraphs:
        eid = id(p._element)
        if eid not in seen:
            seen.add(eid)
            unique.append(p)

    return unique


def _iter_block_items(parent):
    parent_element = parent.element.body if hasattr(parent, "element") and hasattr(parent.element, "body") else parent._element
    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _cell_plain_text(cell) -> str:
    parts = []
    for para in cell.paragraphs:
        text = _paragraph_plain_text(para)
        if text:
            parts.append(text)
    return " | ".join(parts)


def _append_paragraph_once(paragraphs, seen_ids, paragraph, base_contexts, context_hint=""):
    paragraph_id = id(paragraph)
    if paragraph_id in seen_ids:
        return
    seen_ids.add(paragraph_id)
    paragraphs.append(paragraph)
    if context_hint:
        base_contexts[paragraph_id] = context_hint


def _collect_table_paragraphs(table, paragraphs, seen_ids, base_contexts, section_hint=""):
    header_texts = []
    if table.rows:
        header_texts = [_shorten_context_text(_cell_plain_text(cell), 80) for cell in table.rows[0].cells]

    for row_index, row in enumerate(table.rows):
        row_texts = [_shorten_context_text(_cell_plain_text(cell), 80) for cell in row.cells]
        row_anchor = next((text for text in row_texts if text), "")

        for col_index, cell in enumerate(row.cells):
            parts = []
            if section_hint:
                parts.append(f"Section: {section_hint}")

            column_header = ""
            if row_index > 0 and col_index < len(header_texts):
                column_header = header_texts[col_index]
            if column_header:
                parts.append(f"Column: {column_header}")

            if row_anchor and row_anchor != column_header:
                parts.append(f"Row: {row_anchor}")

            context_hint_cell = " | ".join(parts)
            for paragraph in cell.paragraphs:
                _append_paragraph_once(paragraphs, seen_ids, paragraph, base_contexts, context_hint_cell)
            for paragraph in _find_textbox_paragraphs(cell._element, cell):
                _append_paragraph_once(paragraphs, seen_ids, paragraph, base_contexts, context_hint_cell)
            for nested_table in cell.tables:
                _collect_table_paragraphs(nested_table, paragraphs, seen_ids, base_contexts, section_hint)


# ─── Batch Estimation ────────────────────────────────────────

def _estimate_batch_total(texts, contexts, backend, api_key):
    indexed = [(i, t) for i, t in enumerate(texts) if t.strip()]
    if not indexed:
        return 0

    if (not _is_gemini_backend(backend) and backend != "deepl") or not api_key:
        return (len(indexed) + 14) // 15

    if backend == "deepl":
        return (len(indexed) + 7) // 8

    batch_size = _get_gemini_execution_settings(backend)["request_item_limit"]
    return len(_build_adaptive_batches(indexed, contexts or [], batch_size, GEMINI_REQUEST_CHAR_LIMIT))
