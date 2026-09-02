"""
AI Translator — Entry Point
Thin module that wires together core/ modules and exposes translate_document().
All logic lives in core/.
"""

import os
import time
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from core.api_clients import legacy_batch_translate
from core.config import (
    get_api_key, get_deepl_api_key, get_domain_context, get_domain_contexts,
    get_protected_phrases, set_api_key, set_deepl_api_key, set_domain_contexts,
    set_protected_phrases,
)
from core.constants import (
    APP_NAME, FINALIZATION_PROGRESS_STEPS, GEMINI_BACKEND_MODEL_MAP,
    LANGUAGE_CODES, LANGUAGES,
)
from core.docx_parser import (
    _append_paragraph_once, _build_paragraph_contexts, _collect_table_paragraphs,
    _estimate_batch_total, _find_textbox_paragraphs, _is_heading_paragraph,
    _iter_block_items, _legacy_translate_paragraphs, _paragraph_plain_text,
    _prepare_paragraph_units, _translate_paragraphs,
)
from core.progress import TranslationProgress
from core.text_utils import normalize_run_tags
from core.utils import _is_gemini_backend, _summarize_numeric_values


def translate_document(
    input_path: str,
    output_path: str,
    target_lang: str,
    backend: str = "google",
    api_key: str = "",
    progress: TranslationProgress = None,
    domain_context: str = "",
    document_mode: str = "general",
    telemetry: dict | None = None,
    temperature: float | None = None,
) -> str:
    file_ext = Path(input_path).suffix.lower()

    if file_ext == ".pdf":
        from pdf_translator import translate_pdf
        return translate_pdf(
            input_path, output_path, target_lang,
            backend=backend, api_key=api_key,
            progress=progress or TranslationProgress(),
            domain_context=domain_context,
            telemetry=telemetry,
        )

    if file_ext == ".pptx":
        from pptx_translator import translate_presentation
        return translate_presentation(
            input_path, output_path, target_lang,
            backend=backend, api_key=api_key,
            progress=progress or TranslationProgress(),
            domain_context=domain_context,
            telemetry=telemetry,
        )

    if file_ext == ".xlsx":
        from xlsx_translator import translate_spreadsheet
        return translate_spreadsheet(
            input_path, output_path, target_lang,
            backend=backend, api_key=api_key,
            progress=progress or TranslationProgress(),
            domain_context=domain_context,
            telemetry=telemetry,
        )

    if file_ext in (".md", ".markdown"):
        if document_mode != "markdown":
            raise RuntimeError(
                'Select the "Markdown" document mode for .md files. Other modes leave this file type unchanged or misrouted.'
            )
        from md_translator import translate_markdown
        return translate_markdown(
            input_path, output_path, target_lang,
            backend=backend, api_key=api_key,
            progress=progress or TranslationProgress(),
            domain_context=domain_context,
            telemetry=telemetry,
        )

    if file_ext in (".html", ".htm"):
        if document_mode != "html":
            raise RuntimeError(
                'Select the "HTML" document mode for .html files. Other modes leave this file type unchanged or misrouted.'
            )
        from html_translator import translate_html
        return translate_html(
            input_path, output_path, target_lang,
            backend=backend, api_key=api_key,
            progress=progress or TranslationProgress(),
            domain_context=domain_context,
            telemetry=telemetry,
        )

    if document_mode not in {"legacy", "general"}:
        raise RuntimeError(
            f'Document mode "{document_mode}" applies only to Markdown or HTML files. '
            "For .docx, choose General or Legacy."
        )

    if progress is None:
        progress = TranslationProgress()
    progress.reset()
    progress.set_phase("Preparing document...")

    job_started_at = time.perf_counter()
    phrases = get_protected_phrases()
    telemetry = telemetry if telemetry is not None else {}
    telemetry.setdefault("timings", {})
    telemetry.setdefault("document", {})
    load_started_at = time.perf_counter()
    doc = Document(input_path)
    telemetry["timings"]["load_document_seconds"] = round(time.perf_counter() - load_started_at, 3)
    document_mode = (document_mode or "general").strip().lower()
    if document_mode not in {"legacy", "general"}:
        document_mode = "general"
    telemetry["document"].update(
        {
            "document_mode": document_mode,
            "backend": backend,
        }
    )

    def collect_general_paragraphs():
        all_paragraphs = []
        seen_ids = set()
        base_contexts = {}
        current_section_hint = ""

        for block in _iter_block_items(doc):
            if isinstance(block, Paragraph):
                _append_paragraph_once(all_paragraphs, seen_ids, block, base_contexts)
                if _is_heading_paragraph(block):
                    current_section_hint = _paragraph_plain_text(block)
            else:
                _collect_table_paragraphs(block, all_paragraphs, seen_ids, base_contexts, current_section_hint)

        for paragraph in _find_textbox_paragraphs(doc.element, doc):
            _append_paragraph_once(all_paragraphs, seen_ids, paragraph, base_contexts)

        for section in doc.sections:
            if not section.header.is_linked_to_previous:
                try:
                    for block in _iter_block_items(section.header):
                        if isinstance(block, Paragraph):
                            _append_paragraph_once(all_paragraphs, seen_ids, block, base_contexts, "Header")
                        else:
                            _collect_table_paragraphs(block, all_paragraphs, seen_ids, base_contexts, "Header")
                    for paragraph in _find_textbox_paragraphs(section.header._element, section.header):
                        _append_paragraph_once(all_paragraphs, seen_ids, paragraph, base_contexts, "Header")
                except Exception:
                    pass
            if not section.footer.is_linked_to_previous:
                try:
                    for block in _iter_block_items(section.footer):
                        if isinstance(block, Paragraph):
                            _append_paragraph_once(all_paragraphs, seen_ids, block, base_contexts, "Footer")
                        else:
                            _collect_table_paragraphs(block, all_paragraphs, seen_ids, base_contexts, "Footer")
                    for paragraph in _find_textbox_paragraphs(section.footer._element, section.footer):
                        _append_paragraph_once(all_paragraphs, seen_ids, paragraph, base_contexts, "Footer")
                except Exception:
                    pass

        return all_paragraphs, base_contexts

    if document_mode == "legacy":
        collect_started_at = time.perf_counter()
        all_paragraphs = []
        all_paragraphs.extend(doc.paragraphs)
        all_paragraphs.extend(_find_textbox_paragraphs(doc.element, doc))

        def collect_from_table(table):
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)
                    all_paragraphs.extend(_find_textbox_paragraphs(cell._element, cell))
                    for nested_table in cell.tables:
                        collect_from_table(nested_table)

        for table in doc.tables:
            collect_from_table(table)

        for section in doc.sections:
            if not section.header.is_linked_to_previous:
                try:
                    all_paragraphs.extend(section.header.paragraphs)
                    all_paragraphs.extend(_find_textbox_paragraphs(section.header._element, section.header))
                    for table in section.header.tables:
                        collect_from_table(table)
                except Exception:
                    pass
            if not section.footer.is_linked_to_previous:
                try:
                    all_paragraphs.extend(section.footer.paragraphs)
                    all_paragraphs.extend(_find_textbox_paragraphs(section.footer._element, section.footer))
                    for table in section.footer.tables:
                        collect_from_table(table)
                except Exception:
                    pass

        telemetry["timings"]["collect_legacy_seconds"] = round(time.perf_counter() - collect_started_at, 3)
        non_empty_legacy = [para for para in all_paragraphs if para.runs and _paragraph_plain_text(para)]
        estimated_batches = 0
        if (not _is_gemini_backend(backend) and backend != "deepl") or not api_key:
            estimated_batches = (len(non_empty_legacy) + 14) // 15
        elif backend == "deepl":
            estimated_batches = (len(non_empty_legacy) + 7) // 8
        elif backend in {"gemini", "gemini-25-flash", "gemini-37-flash"}:
            estimated_batches = (len(non_empty_legacy) + 11) // 12
        elif backend == "gemini-35-flash-lite":
            estimated_batches = (len(non_empty_legacy) + 11) // 12
        else:
            estimated_batches = (len(non_empty_legacy) + 14) // 15
        total_steps = estimated_batches + FINALIZATION_PROGRESS_STEPS
        progress.set_total(max(1, total_steps))
        telemetry["document"].update(
            {
                "legacy_paragraph_count": len(all_paragraphs),
                "legacy_non_empty_paragraph_count": len(non_empty_legacy),
                "estimated_total_steps": total_steps,
            }
        )

        progress.set_phase("Translating document...")
        translation_started_at = time.perf_counter()
        _legacy_translate_paragraphs(
            all_paragraphs,
            target_lang,
            backend,
            api_key,
            phrases,
            progress,
            domain_context,
            telemetry=telemetry,
        )
        telemetry["timings"]["translate_legacy_seconds"] = round(time.perf_counter() - translation_started_at, 3)
    else:  # general mode
        collect_started_at = time.perf_counter()
        all_paragraphs, base_contexts = collect_general_paragraphs()
        telemetry["timings"]["collect_general_seconds"] = round(time.perf_counter() - collect_started_at, 3)
        telemetry["document"].update(
            {
                "general_paragraph_count": len(all_paragraphs),
                "general_base_context_count": len(base_contexts),
            }
        )

        progress.set_phase("Analyzing context...")
        context_started_at = time.perf_counter()
        paragraph_contexts = _build_paragraph_contexts(all_paragraphs, base_contexts)
        telemetry["timings"]["build_context_seconds"] = round(time.perf_counter() - context_started_at, 3)
        prepared_units = _prepare_paragraph_units(all_paragraphs, paragraph_contexts)
        general_payloads = [unit["fused_text"] for unit in prepared_units]
        general_payload_contexts = [unit.get("context", "") for unit in prepared_units]
        total_steps = _estimate_batch_total(general_payloads, general_payload_contexts, backend, api_key) + FINALIZATION_PROGRESS_STEPS
        progress.set_total(max(1, total_steps))
        telemetry["document"].update(
            {
                "general_unit_count": len(prepared_units),
                "general_payload_char_stats": _summarize_numeric_values([len(payload) for payload in general_payloads]),
                "estimated_total_steps": total_steps,
            }
        )

        progress.set_phase("Translating document...")
        translation_started_at = time.perf_counter()
        _translate_paragraphs(
            all_paragraphs,
            target_lang,
            backend,
            api_key,
            phrases,
            progress,
            domain_context,
            paragraph_contexts=paragraph_contexts,
            telemetry=telemetry,
            set_total=False,
            temperature=temperature,
        )
        telemetry["timings"]["translate_general_seconds"] = round(time.perf_counter() - translation_started_at, 3)

    progress.set_phase("Saving document...")
    save_started_at = time.perf_counter()
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    telemetry["timings"]["save_document_seconds"] = round(time.perf_counter() - save_started_at, 3)
    telemetry["timings"]["total_job_seconds"] = round(time.perf_counter() - job_started_at, 3)
    progress.advance("Translation complete")

    return output_path
